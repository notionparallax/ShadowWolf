
# In[5]:

from docx import Document
import os
import pandas as pd


# In[11]:

docs =  os.listdir("./Desktop/projectText/")

docs.remove("110704__Sulman_Award_list.docx") #this doc has a left column full of dates
docs.remove("131021__PlantBank_building_info.docx")#this doc is an article
#docs# = docs[:10]


# In[12]:

def isShaded(aRun):
    
    try:
        xml = aRun._r.xml.encode('ascii', 'replace')
        #print xml
        if u"<w:shd w:fill" in xml or u"<w:highlight" in xml:            
            return aRun.text
        else:
            return ""
    except UnicodeEncodeError:
         return "MC"


# In[14]:

rows = []
for fileName in docs:
    row = {"filename": fileName}
    doc = Document("./Desktop/projectText/" + fileName)
    #print fileName, len(doc.tables)
    if len(doc.tables) > 0:
        for table in doc.tables:
                for trow in table.rows:
                    if len(trow.cells) == 2:
                        key   = ""
                        value = ""
                        protoKey = ""
                        for p in trow.cells[0].paragraphs:
                            protoKey += p.text

                        text = protoKey

                        textin = text

                        text = text.strip()              #remove padding whitespace

                        text = text.upper()             #make it all uppercase                        

                        text = text.replace(u":"," ")      #remove colons
                        text = text.replace(u"\u2013","-") #remove em dashes
                        text = text.replace(u"\u2019","'") #remove pretty '
                        text = text.replace(u"/"," / ")    #space either side of a /
                        text = text.replace(u"-"," - ")    #space either side of a -
                        text = text.replace(u"  "," ")     #remove double spaces
                        text = text.replace(u"  "," ")     #remove double spaces
                        text = text.replace(u"CAN CHOOSE MULTIPLE","")
                        text = text.replace(u"CAN CHOOSE MUTIPLE","")
                        text = text.replace(u"SHEETDESCRIPTION","SHEET DESCRIPTION")#remove double spaces
                        text = text.replace(u"DESCRIPTIONTEXT","DESCRIPTION TEXT")  #remove double spaces
                        text = text.replace(u"FUNK","FUNC")
                        text = text.replace(u"BVNDH","BVN") #just to keep it consistent (sorry BDon)
                        text = text.replace(u"?","")
                        text = text.replace(u"MASTERPLAN OPTION","")
                        text = text.replace(u"DESUBG","DESIGN")      #seems like a typo
                        text = text.replace(u"REFEREECURRENT","REFEREE CURRENT")
                        text = text.replace(u"E.G. CONFIDENTIAL","") 
                        text = text.replace(u"PLEASE REFER TO THE COMMUNICATIONS MANUAL FOR IMAGE REQUIREMENTS","")
                        text = text.replace(u"PLEASE TICK.","")


                        text = text.split("(")[0]        #remove anything after a bracket
                        text = text.split("INCLUDING")[0]
                        text = text.split("*")[0]
                        text = text.split("/ GBCA")[0]        #remove anything after a bracket
                        text = text.split("WHAT IS ITWHAT DOES IT AIM TO DO")[0]
                        text = text.split("FROM PRINCIPAL SIGNED-OFF CREDIT LIST")[0]
                        text = text.split("ANY INNOVATIVE ASPECTS")[0]
                        text = text.split("MOST RELEVANT FOR HEALTH")[0]
                        text = text.split(")")[0]
                        text = text.split("MUST USE THIS FOR")[0]
                        text = text.split("BOLD =")[0]
                        text = text.split("AS OF 9")[0]
                        text = text.split("/ S")[0]
                        if "PROJECT SHEETS" in text:
                            text = text.split("PROJECT SHEETS")[0] + "PROJECT SHEETS"

                        text = text.strip()              #remove padding whitespace
                        #print u"|{}| --> |{}|".format(textin,text)
                        key   += u"{}".format(text)

                        for p in trow.cells[1].paragraphs:
                            if "CATEGOR" in key:
                                #print fileName, key
                                for r in p.runs:
                                    #print r.text
                                    shadedCat = isShaded(r).strip()
                                    if len(shadedCat) > 0:
                                        value += u"{},".format(shadedCat)
                                    
                                #print value
                            else:
                                value += u"<p>{}</p>".format(p.text.strip())
                        
                        row[key]= value
                        row["numCols"] = 2

                    else:
                        row["numCols"] = len(trow.cells)
    rows.append(row)


# In[15]:

projects = pd.DataFrame.from_dict(rows)
projects.head()


# In[16]:

duffCols = [u"PROMOTIONAL SUMMARYTO BE USED FOR THE INSTITUTE'S WEBSITE AND IN AWARDS MATERIALS",
            u'ABW AT ASB NORTH WHARF',
            u'ADDITIONALPROJECT SHEETUNIVERSITY OF ADELAIDE -THE BRAGGS,-MAWSON LABORATORIES-SCOTT THEATRE-FLOREY HEALTH SCIENCES HUB',
            u'ARUP',
            u'ASB NORTH WHARF - LEGACY',
            u'CISCO SYSTEMS',
            u'CONTINUED DISCUSSION RE POST OCCUPANCY EVALUATION',
            u'EXAMPLES OF RESEARCH PROJECTS UNDERTAKEN AT NALSH IN 2013',
            u'SALUTOGENIC APPROACH',
            u'STEPHANIE BRINCAT',
            u'UMOW LAI',
            u'WILKINSON MURRAY',
            u'WT PARTNERSHIP',
            u'WEBSITE UPDATE AECOM WORKPLACE SYDNEY',
            u'WEBSITE UPDATE AUSTRALIAN RED CROSS BLOOD SERVICES NSW HEADQUARTERS',
            u'WEBSITE UPDATE ELIZABETH STREET TOWER',
            u'SUBJECTFAXEDPOSTEDNUMBER OF PAGES',
            u'TO',u"TTW",u"CC", u'CUNDALLS',u"COCKRAM CONSTRUCTION",
            u"FROM",u"SANDRICK",u"NJM DESIGN",
            u"ELIZABETH HUGHES",u"IRELAND BROWN CONSTRUCTIONS",
            
            #single value columns - These will need to be manually investigated later on
            u"/ VALUE OUTCOME",
            u"ALL STAFF WORKSHOP",
            u"ARCHITECTURAL EXPRESSION OF CONCEPT",
            u"ART WORK",
            u"AUDIO VISUAL",
            u"BRIEF ELEMENTS - HOSPITALS / LABS",
            u"BRIEFING",
            u"BRIEFING PROCESS - 150 WORDS",
            u"BUDGET",
            u"BVN CONTACT",
            u"CATEGORY - SUB CATEGORY",
            u"CLIENT - USER - BUILDING OWNER",
            u"CLIENT CONCERNS & RESOLUTIONS",
            u"CLIENT CONCERNS AND RESOLUTIONS",
            u"CLIENT NAME",
            u"CLIENT TESTIMONIALFOR STANDARD PROJECT SHEET",
            u"COLLABORATION",
            u"COMMUNITY",
            u"COMPETITION",
            u"COMPETITION STAGE",
            u"COMPLETION",
            u"CONCEPTUAL FRAMEWORK",
            u"DATE",
            u"DATE COMPLETED",
            u"DATES",
            u"DAVIS LANGDON",
            u"DELIVERY DETAILS",
            u"DESCRIPTION - PROJECT SHEET - PROPOSED TEXT FOR PROJECT SHEETS",
            u"DESCRIPTION - PROJECT SHEET200 WORDS",
            u"DESCRIPTION - PROJECT SHEETMACARTHUR CAMPBELLTOWN HOSPITAL MASTERPLAN",
            u"DESIGN DETAILS",
            u"DESIGN INTENTIONS",
            u"DESIGN PRINCIPLES",
            u"DESIGN STATEMENT",
            u"ENGAGEMENT WITH PUBLIC REALM",
            u"ENVIRONMENTAL, SOCIAL AND ECONOMIC SUSTAINABILITY",
            u"FACULTY WORKSHOP",
            u"FULL DEPARTMENTAL LIST - BRIEFED ELEMENTS",
            u"FULL DEPARTMENTAL LIST OF BRIEF ELEMENTS",
            u"FUNCTION",
            u"HERITAGE",
            u"INCLUSIVE DESIGN",
            u"INTEGRATION OF ALLIED DISCIPLINES",
            u"INTERIOR RESPONSE",
            u"KEY BUILDING FACTS",
            u"KEY DESIGN CONCEPT",
            u"LATITUDE / LONGITUDE",
            u"LEARNING ENVIRONMENT",
            u"MAIN CONCEPT",
            u"MATERIAL & DETAILING",
            u"MEMORABLE MOMENT",
            u"OFFICE",
            u"OFFIDESCRIPTION - PROJECT SHEETS",
            u"ORGANISATION OF SPACES & FUNCTIONS",
            u"OWNER / CLIENT CONTACT",
            u"PHYSICAL ENVIRONMENT",
            u"PLANNING PROCESS",
            u"PROGRAM RESOLUTION",
            u"PROJECT CATEGORY",
            u"PROJECT PHILOSOPHY / METHODOLOGY",
            u"PROJECT RELEVANCE",
            u"PROJECT SHEET DESCRIPTION",
            u"PROJECT SUBCATEGORIES",
            u"PROJECTION LIGHTING",
            u"PUBLIC CULTURAL BENEFITS",
            u"PUBLIC REALM ENGAGEMENT",
            u"REFEREE CURRENT",
            u"REFURBISHMENT",
            u"RELATIONSHIP OF BUILDING TO ITS CONTEXT",
            u"RELATIONSHIP OF BUILT FORM TO CONTEXT",
            u"RESPONSE TO CLIENT USER NEEDS",
            u"RETAIL COMPONENT",
            u"ROLE",
            u"SCOPE",
            u"SERVICE PROVIDED",
            u"SERVICES CONSULTANT",
            u"SIZEGFA / NLA M2",
            u"STAKEHOLDER CONSULTATION",
            u"STAKEHOLDERS USER GROUPS WITH STUDENTS AND STAFF",
            u"STUDENT WORKSHOP",
            u"SUBJECT",
            u"SUBMISSION PROCESS",
            u"SUSTAINABILITY",
            u"TECHNOLOGY",
            u"URBAN DESIGN",
            u"VALUE",
            u"VALUE ADD TO THE BUSINESS"
            #u'WILKINSON MURRAY',
            #u'WT PARTNERSHIP'
            ]

colList = projects.columns.tolist()

for duffCol in duffCols:
    if duffCol in colList:
        projects = projects.drop(duffCol, axis=1)
        #print "DROP WIN!!! " + duffCol 
    else:
        print "couldn't drop " + duffCol 


# In[25]:

projects.to_excel("./projects.xls")
projects.head()


# In[26]:

projects.columns.tolist()


# In[27]:

for col in projects.columns.tolist():
    valList=projects[col].tolist()
    print u"{:03} : {}".format(len(valList) - valList.count(NaN), col)







# In[ ]:


